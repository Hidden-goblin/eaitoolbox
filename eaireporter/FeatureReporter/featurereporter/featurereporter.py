# -*- coding: utf-8 -*-
# -*- Product under GNU GPL v3 -*-
import re
import glob
import logging
import argparse
import os
import sys

from pathlib import Path
from PIL import ImageTk, Image
from behave.parser import parse_file
from docx import Document
import tkinter as tk
from tkinter import filedialog, Toplevel, messagebox

log = logging.getLogger(__name__)

LICENCE = """ ExportUtilities  Copyright (C) 2021  E.Aivayan
    This program comes with ABSOLUTELY NO WARRANTY.
    This is free software, and you are welcome to redistribute it under certain conditions.
    
    Please see https://opensource.org/licenses/GPL-3.0
    """

class Application:
    def __init__(self):
        self.__assets = os.path.dirname(os.path.realpath(__file__))
        self.__master = tk.Tk()
        self.__master.geometry("500x200")
        # Feature repository vars
        self.__repository_label = None
        self.__repository_select_button = None
        self.__repository_location = None
        self.__respository_status = None
        self.__picture_valid = None
        self.__picture_warning = None
        # Created document
        self.__document_name_label = None
        self.__document_name_input = None
        self.__document_filename_label = None
        self.__document_filename_input = None
        self.__us_tag_label = None
        self.__us_tag_input = None
        # Execution reference
        self.__execution_result_label = None
        self.__execution_result_status = None
        self.__execution_result_button = None
        self.__excution_location = None
        # Other UI thing
        self.__quit = None
        self.__readme_button = None
        self.__execute_button = None
        self.__legal_label = None
        # Reporter object
        self.__reporter = ExportUtilities()
        # Create
        self.create_widgets()
        self.create_layout()



    def create_widgets(self):
        # Legal stuff
        self.__legal_label = tk.Label(self.__master, text="?", )
        self.__legal_label.bind("<Button-1>", self.__display_legal)
        # Picture stuff
        valid = Image.open(f"{self.__assets}/assets/valid.png")
        valid = valid.resize((20, 20), Image.ANTIALIAS)
        self.__picture_valid = ImageTk.PhotoImage(valid)
        warning = Image.open(f"{self.__assets}/assets/warning.png")
        warning = warning.resize((20, 20), Image.ANTIALIAS)
        self.__picture_warning = ImageTk.PhotoImage(warning)
        # Repository
        self.__respository_status = tk.Label(self.__master, image=self.__picture_warning)
        self.__repository_label = tk.Label(self.__master,
                                           text="Please select a feature file repository.",
                                           wraplength="250")
        self.__repository_select_button = tk.Button(self.__master,
                                                    text="Select repository",
                                                    command=self.__select_repository)
        # Document name
        self.__document_name_label = tk.Label(self.__master,
                                              text="Document name: ")
        self.__document_name_input = tk.Entry(self.__master)
        # Document filename
        self.__document_filename_label = tk.Label(self.__master,
                                                  text="Document filname: ")
        self.__document_filename_input = tk.Entry(self.__master)
        # US Tag
        self.__us_tag_label = tk.Label(self.__master,
                                       text="US Tag: ")
        self.__us_tag_input = tk.Entry(self.__master)
        # Execution location
        self.__execution_result_label = tk.Label(self.__master,
                                                 text="Execution results location: ")
        self.__execution_result_button = tk.Button(self.__master, text="Select execution",
                                                   command=self.__select_execution)
        self.__execution_result_reset = tk.Button(self.__master, text="Reset location",
                                                  command=self.__reset_execution)
        self.__execution_result_status = tk.Label(self.__master, text="")
        # Readme
        self.__readme_button = tk.Button(self.__master, text="README",
                                         command=self.__display_readme)

        # Execute
        self.__execute_button = tk.Button(self.__master, text="Create report",
                                          command=self.__create_report)

        # QUIT
        self.__quit = tk.Button(self.__master, text="QUIT", fg="red",
                              command=self.__master.destroy)

    def create_layout(self):
        self.__legal_label.grid(row=0, column=4)
        self.__respository_status.grid(row=1, column=0)
        self.__repository_label.grid(row=1, column=1)
        self.__repository_select_button.grid(row=1, column=3, columnspan=4)
        self.__document_name_label.grid(row=2, column=0)
        self.__document_name_input.grid(row=2, column=1)
        self.__document_filename_label.grid(row=3, column=0)
        self.__document_filename_input.grid(row=3, column=1)
        self.__us_tag_label.grid(row=4, column=0)
        self.__us_tag_input.grid(row=4, column=1)
        self.__execution_result_label.grid(row=5, column=0)
        self.__execution_result_status.grid(row=5, column=1)
        self.__execution_result_button.grid(row=5, column=3)
        self.__execution_result_reset.grid(row=5, column=4)
        self.__readme_button.grid(row=6, column=0)
        self.__execute_button.grid(row=6, column=3)
        self.__quit.grid(row=7, column=0, columnspan=5, sticky="E,W")

    def __display_readme(self):
        messagebox.showinfo("Quick manual",
                            """1- Select the folder where you store the feature files
 Optionaly:
     2- Select the report title
     3- Select the report file name
     4- Select the tag linking to US
     5- Select the behave plain report file""")

    def __create_report(self):
        if self.__repository_location is not None and self.__repository_location:
            self.__reporter.feature_repository = self.__repository_location
            if self.__document_name_input.get():
                self.__reporter.report_title = self.__document_name_input.get()
            if self.__us_tag_input.get():
                self.__reporter.us_tag = self.__us_tag_input.get()
            param = dict()
            if self.__document_filename_input.get():
                param["output_file_name"] = self.__document_filename_input.get()
            if self.__excution_location is not None and self.__excution_location:
                param["report_file"] = self.__excution_location
            print(param)
            self.__reporter.create_application_documentation(**param)
        else:
            messagebox.showerror("Report creation", "Cannot create de report without a feature files repository.\n Please select one.")

    def __display_legal(self, event):
        print("Display legal")
        fInfos = Toplevel()  # Popup -> Toplevel()
        fInfos.title('Infos')
        text = tk.Text(fInfos, height=15, width=90)
        text.insert(tk.END,
                    f"""License
*******

{LICENCE}

Pictures disclaimer
*******************
 
Icon by Raj Dev (https://freeicons.io/profile/714) on https://freeicons.io""")
        text.grid(row=0, column=0)
        tk.Button(fInfos, text='Quitter', command=fInfos.destroy).grid(row=1, column=0)
        fInfos.transient(self.__master)  # Réduction popup impossible
        fInfos.grab_set()  # Interaction avec fenetre jeu impossible
        self.__master.wait_window(fInfos)  # Arrêt script principal

    def __select_repository(self):
        self.__repository_location = filedialog.askdirectory(parent=self.__master, mustexist=True, title="Select the feature repository")
        if self.__repository_location is not None and self.__repository_location:
            self.__repository_label["text"] = self.__repository_location
            self.__respository_status["image"] = self.__picture_valid
        else:
            self.__repository_label["text"] = "Please select a feature file repository."
            self.__respository_status["image"] = self.__picture_warning
        # self.__respository_status.configure(image=self.__picture_valid)
        # self.__respository_status.image = self.__picture_valid

    def __select_execution(self):
        self.__excution_location = filedialog.askopenfilename(parent=self.__master,
                                                              title="Select the test plain report",
                                                              filetypes=[ ("text files", "*.txt") ])
        if self.__excution_location is not None and self.__excution_location:
            self.__execution_result_status["text"] = "Execution selected"
        else:
            self.__execution_result_status["text"] = ""

    def __reset_execution(self):
        self.__execution_result_status["text"] = ""
        self.__excution_location = None

    def run(self):
        self.__master.mainloop()



class ExportUtilities:

    def __init__(self, feature_repository: str = None,
                 user_story_tag_prefix: str = None,
                 report_title: str = None):
        self.__feature_repository = feature_repository
        self.__user_story_tag_prefix = user_story_tag_prefix
        self.__report_title = report_title
        self.__document = None

    @property
    def feature_repository(self):
        return self.__feature_repository

    @feature_repository.setter
    def feature_repository(self, repository: str):
        if Path(repository).exists() and Path(repository).is_dir():
            self.__feature_repository = repository
        else:
            raise FileExistsError(f"{repository} is not a existing folder")

    @property
    def report_title(self):
        return self.__report_title

    @report_title.setter
    def report_title(self, title: str):
        if isinstance(title, str) and title:
            self.__report_title = title
        else:
            raise AttributeError(f"{title} must be a non empty string")

    @property
    def us_tag(self):
        return self.__user_story_tag_prefix

    @us_tag.setter
    def us_tag(self, tag_value: str):
        if isinstance(tag_value, str) and tag_value:
            self.__user_story_tag_prefix = tag_value
        else:
            log.warning("Use the None value")
            self.__user_story_tag_prefix = None

    @property
    def document(self):
        return self.__document

    def create_application_documentation(self, report_file=None, output_file_name="demo.docx"):
        """
        Create a document (docx) object and read first all ".feature" files and
        add their contents into the document.

        If a report file name is provided, it will add a "last execution" section containing
        the data found in the report.

        The report file must be the "plain" report output file generated from behave.

        :param report_file: The report file path (absolute or relative)
        :param output_file_name : The exported file name by default "demo.docx"
        :return: None
        """

        self.__document = Document()
        self.document.add_heading("{}".format(self.__report_title), 0)  # Document title
        self.document.add_page_break()
        for file in glob.iglob("{}/**/*.feature".format(self.__feature_repository),
                               recursive=True):  # Use the iterator as it's cleaner
            # log.debug()
            test = parse_file(file)  # Use the Behave parser in order to read the feature file
            self.add_heading(feature=test)
            self.add_description(feature=test)
            self.add_background(feature=test)
            self.add_scenario(feature=test)
            self.document.add_page_break()
        if report_file is not None:
            self.add_report(file=report_file)
        self.document.save(output_file_name)

    def add_heading(self, feature=None):
        """
        Add a the feature name as top level section
        :param feature: the feature object
        :return:
        """
        self.document.add_heading(feature.name, 1)
        paragraph = self.document.add_paragraph("")
        if self.us_tag is not None:
            matcher = [elem for elem in feature.tags if self.us_tag in elem]
            paragraph.add_run("Related to the user story: {}".format(str(matcher).strip('[]')))

    def add_description(self, feature=None):
        """
        Add the feature description into the document.
        If the feature file line contains a "*" character it will create a bullet list.
        If the line contains the sentence "business rules" then it will print it in bold.
        Otherwise it will print the line as-is.
        :param feature: the feature object
        :return: None
        """
        for line in feature.description:
            if re.match(r'\*.*', line):
                self.document.add_paragraph(line[1:], style='List Bullet')
            elif re.match('[Bb]usiness [Rr]ules.*', line):
                paragraph = self.document.add_paragraph("")
                paragraph.add_run(line).bold = True
            else:
                self.document.add_paragraph(line)

    def add_background(self, feature=None):
        """
        Add the background of each scenario into the document
        :param feature: the feature object from where to retrieve the background
        :return: None
        """
        if feature.background is not None:
            self.print_scenario_title(scenario_keyword=feature.background.keyword,
                                      scenario_name=feature.background.name)
            self.print_steps(steps=feature.background.steps)

    def add_scenario(self, feature=None):
        """
        Add in the document all the scenarios attached to a feature.
        :param feature: the feature object
        :return: None
        """
        if feature.scenarios is not None:
            for scenario in feature.scenarios:
                self.print_scenario_title(scenario_keyword=scenario.keyword,
                                          scenario_name=scenario.name)
                self.print_steps(steps=scenario.steps)
                if scenario.type == 'scenario_outline':
                    self.print_examples(examples=scenario.examples)

    def print_examples(self, examples=None):
        """
        Add an example section for each example attached to a scenario outline
        :param examples: an example list
        :return: None
        """
        for example in examples:
            self.print_scenario_title(scenario_keyword=example.keyword,
                                      scenario_name=example.name, level=3)
            self.print_table(table=example.table)

    def print_scenario_title(self, scenario_keyword=None, scenario_name=None, level=2):
        """
        Print a section title in the document in the format keyword : name with a level 2
        :param scenario_keyword: the keyword such as Scenario or Scenario Outline or Example
        :param scenario_name: the scenario name
        :param level: the section level. By default level 2
        :return: None
        """
        self.document.add_heading(f"{scenario_keyword}: {scenario_name}", level=level)
        # paragraph = document.add_paragraph("")
        # paragraph.add_run("{}:".format(scenario_keyword)).bold = True
        # paragraph.add_run(scenario_name)

    def print_steps(self, steps=None):
        """
        Add a step section into the document.
        :param steps: the feature steps table
        :return: None
        """
        step_done = []
        for step in steps:
            if step.keyword in step_done:
                keyword = "And"
            else:
                step_done.append(step.keyword)
                keyword = step.keyword

            paragraph = self.document.add_paragraph("", style='No Spacing')
            paragraph.add_run(keyword).bold = True
            paragraph.add_run(" {}".format(step.name))

            if step.table is not None:
                self.print_table(table=step.table)

    def print_table(self, table=None):
        """
        Add a tabular word object in the document based on the feature table object
        :param table: a feature table object
        :return: None
        """
        number_of_column = len(table.headings)

        table_instance = self.document.add_table(rows=1,
                                                 cols=number_of_column,
                                                 style='Light List Accent 3')
        header_cells = table_instance.rows[0].cells
        for count, text in enumerate(table.headings):
            header_cells[count].text = str(text)

        for row in table.rows:
            row_cells = table_instance.add_row().cells
            for count, cell in enumerate(row.cells):
                row_cells[count].text = str(cell)

    def add_report(self, file=None):
        """
        Add a last execution section to a document. It reads an execution plain file report
        :param file: the file name (relative or absolute) of the execution report.
        :return: None
        """
        self.document.add_heading("Last Execution report", 1)
        reporter = {}
        failed_count = 0
        succeed_count = 0
        test_count = 0
        current_feature = None
        current_scenario = None
        last_status = "skipped"
        with open(file) as report_file:
            for line in report_file.readlines():
                if re.match("Feature.*", line):
                    if current_feature is not None:
                        reporter[current_feature].update({current_scenario: last_status})
                        if last_status == "failed":
                            failed_count += 1
                        elif last_status == "passed":
                            succeed_count += 1
                        test_count += 1
                        last_status = "skipped"
                        self.document.add_page_break()
                    current_feature = line.split(":")[1].lstrip(' ').rstrip()
                    reporter[current_feature] = {}
                    current_scenario = None  # No scenario for the current feature
                    print(reporter)
                    self.document.add_heading(line.rstrip(), 2)
                elif re.match(r"\s*Scenario.*", line):
                    if current_scenario is not None:
                        reporter[current_feature].update({current_scenario: last_status})
                        if last_status == "failed":
                            failed_count += 1
                        elif last_status == "passed":
                            succeed_count += 1
                        test_count += 1
                        last_status = "skipped"
                    current_scenario = line.split(":")[1].lstrip(' ').rstrip()
                    print(current_scenario)
                    self.document.add_heading(line.rstrip(), 3)
                elif re.match(".*passed.*", line):
                    last_status = "passed"
                    self.document.add_paragraph(line.rstrip(), style='No Spacing')
                elif re.match(".*failed.*", line):
                    last_status = "failed"
                    self.document.add_paragraph(line.rstrip(), style='No Spacing')
                else:
                    self.document.add_paragraph(line.rstrip(), style='No Spacing')

        self.document.add_page_break()
        self.document.add_heading("Last Execution summary", 1)
        table_instance = self.document.add_table(rows=1, cols=3, style='Light List Accent 3')
        header_cells = table_instance.rows[0].cells
        header_cells[0].text = "Feature"
        header_cells[1].text = "Scenario"
        header_cells[2].text = "Status"
        print(reporter)
        feature_keys = list(reporter.keys())
        feature_keys.sort()
        for feature_key in feature_keys:
            scenario_keys = list(reporter[feature_key].keys())
            scenario_keys.sort()
            print(scenario_keys)
            for scenario_key in scenario_keys:
                row_cells = table_instance.add_row().cells
                row_cells[0].text = feature_key
                row_cells[1].text = scenario_key
                row_cells[2].text = reporter[feature_key][scenario_key]


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--tag", help="Invariant pointing to a user story")
    parser.add_argument("--title", help="The document's title")
    parser.add_argument("--repository", help="The folder where the feature files are")
    parser.add_argument("--output", help="")
    parser.add_argument("--execution",
                        help="Behave plain test output in order to also print the last execution result")
    parser.add_argument("--license",
                        help="Display the license.",
                        action="store_true")

    args = parser.parse_args()
    if all([value is None for item, value in vars(args).items() if item != "license"]) and not args.license:
        app = Application()
        app.run()
    else:
        print(args.license)
        if args.license is not None and args.license:
            with open(os.path.realpath(f"{os.path.dirname(os.path.realpath(__file__))}/assets/LICENSE.txt")) as license:
                print(license.read())
                sys.exit(0)
        if args.repository is None or not args.repository:
            parser.print_help()
        report = ExportUtilities()
        report.feature_repository = args.repository
        if args.title is not None and args.title:
            report.report_title = args.title
        if args.tag is not None and args.tag:
            report.us_tag = args.tag
        parameters = dict()
        if args.execution is not None and args.execution:
            parameters["report_file"] = args.execution
        if args.output is not None and args.output:
            parameters["output_file_name"] = args.output
        print(f"""{LICENCE}
    Run with --license option to display the full licence""")
        report.create_application_documentation(**parameters)
    sys.exit(0)

if __name__ == '__main__':
    main()