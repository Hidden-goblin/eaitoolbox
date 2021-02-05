# -*- coding: utf-8 -*-
import re
import glob
import logging
import argparse

from pathlib import Path
from behave.parser import parse_file
from docx import Document

log = logging.getLogger(__name__)


class ExportUtilities:

    def __init__(self, feature_repository: str = None,
                 user_story_tag_prefix: str = None,
                 report_title: str = None):
        self.__feature_repository = feature_repository
        self.__user_story_tag_prefix = user_story_tag_prefix
        self.__report_title = report_title
        self.__document = None

    @property
    def feature_repository(self):
        return self.__feature_repository

    @feature_repository.setter
    def feature_repository(self, repository: str):
        if Path(repository).exists() and Path(repository).is_dir():
            self.__feature_repository = repository
        else:
            raise FileExistsError(f"{repository} is not a existing folder")

    @property
    def report_title(self):
        return self.__report_title

    @report_title.setter
    def report_title(self, title: str):
        if isinstance(title, str) and title:
            self.__report_title = title
        else:
            raise AttributeError(f"{title} must be a non empty string")

    @property
    def us_tag(self):
        return self.__user_story_tag_prefix

    @us_tag.setter
    def us_tag(self, tag_value: str):
        if isinstance(tag_value, str) and tag_value:
            self.__user_story_tag_prefix = tag_value
        else:
            log.warning("Use the None value")
            self.__user_story_tag_prefix = None

    @property
    def document(self):
        return self.__document

    def create_application_documentation(self, report_file=None, output_file_name="demo.docx"):
        """
        Create a document (docx) object and read first all ".feature" files and
        add their contents into the document.

        If a report file name is provided, it will add a "last execution" section containing
        the data found in the report.

        The report file must be the "plain" report output file generated from behave.

        :param report_file: The report file path (absolute or relative)
        :param output_file_name : The exported file name by default "demo.docx"
        :return: None
        """

        self.__document = Document()
        self.document.add_heading("{}".format(self.__report_title), 0)  # Document title
        self.document.add_page_break()
        for file in glob.iglob("{}/**/*.feature".format(self.__feature_repository),
                               recursive=True):  # Use the iterator as it's cleaner
            # log.debug()
            test = parse_file(file)  # Use the Behave parser in order to read the feature file
            self.add_heading(feature=test)
            self.add_description(feature=test)
            self.add_background(feature=test)
            self.add_scenario(feature=test)
            self.document.add_page_break()
        if report_file is not None:
            self.add_report(file=report_file)
        self.document.save(output_file_name)

    def add_heading(self, feature=None):
        """
        Add a the feature name as top level section
        :param feature: the feature object
        :return:
        """
        self.document.add_heading(feature.name, 1)
        paragraph = self.document.add_paragraph("")
        if self.us_tag is not None:
            matcher = [elem for elem in feature.tags if self.us_tag in elem]
            paragraph.add_run("Related to the user story: {}".format(str(matcher).strip('[]')))

    def add_description(self, feature=None):
        """
        Add the feature description into the document.
        If the feature file line contains a "*" character it will create a bullet list.
        If the line contains the sentence "business rules" then it will print it in bold.
        Otherwise it will print the line as-is.
        :param feature: the feature object
        :return: None
        """
        for line in feature.description:
            if re.match(r'\*.*', line):
                self.document.add_paragraph(line[1:], style='List Bullet')
            elif re.match('[Bb]usiness [Rr]ules.*', line):
                paragraph = self.document.add_paragraph("")
                paragraph.add_run(line).bold = True
            else:
                self.document.add_paragraph(line)

    def add_background(self, feature=None):
        """
        Add the background of each scenario into the document
        :param feature: the feature object from where to retrieve the background
        :return: None
        """
        if feature.background is not None:
            self.print_scenario_title(scenario_keyword=feature.background.keyword,
                                      scenario_name=feature.background.name)
            self.print_steps(steps=feature.background.steps)

    def add_scenario(self, feature=None):
        """
        Add in the document all the scenarios attached to a feature.
        :param feature: the feature object
        :return: None
        """
        if feature.scenarios is not None:
            for scenario in feature.scenarios:
                self.print_scenario_title(scenario_keyword=scenario.keyword,
                                          scenario_name=scenario.name)
                self.print_steps(steps=scenario.steps)
                if scenario.type == 'scenario_outline':
                    self.print_examples(examples=scenario.examples)

    def print_examples(self, examples=None):
        """
        Add an example section for each example attached to a scenario outline
        :param examples: an example list
        :return: None
        """
        for example in examples:
            self.print_scenario_title(scenario_keyword=example.keyword,
                                      scenario_name=example.name, level=3)
            self.print_table(table=example.table)

    def print_scenario_title(self, scenario_keyword=None, scenario_name=None, level=2):
        """
        Print a section title in the document in the format keyword : name with a level 2
        :param scenario_keyword: the keyword such as Scenario or Scenario Outline or Example
        :param scenario_name: the scenario name
        :param level: the section level. By default level 2
        :return: None
        """
        self.document.add_heading(f"{scenario_keyword}: {scenario_name}", level=level)
        # paragraph = document.add_paragraph("")
        # paragraph.add_run("{}:".format(scenario_keyword)).bold = True
        # paragraph.add_run(scenario_name)

    def print_steps(self, steps=None):
        """
        Add a step section into the document.
        :param steps: the feature steps table
        :return: None
        """
        step_done = []
        for step in steps:
            if step.keyword in step_done:
                keyword = "And"
            else:
                step_done.append(step.keyword)
                keyword = step.keyword

            paragraph = self.document.add_paragraph("", style='No Spacing')
            paragraph.add_run(keyword).bold = True
            paragraph.add_run(" {}".format(step.name))

            if step.table is not None:
                self.print_table(table=step.table)

    def print_table(self, table=None):
        """
        Add a tabular word object in the document based on the feature table object
        :param table: a feature table object
        :return: None
        """
        number_of_column = len(table.headings)

        table_instance = self.document.add_table(rows=1,
                                                 cols=number_of_column,
                                                 style='Light List Accent 3')
        header_cells = table_instance.rows[0].cells
        for count, text in enumerate(table.headings):
            header_cells[count].text = str(text)

        for row in table.rows:
            row_cells = table_instance.add_row().cells
            for count, cell in enumerate(row.cells):
                row_cells[count].text = str(cell)

    def add_report(self, file=None):
        """
        Add a last execution section to a document. It reads an execution plain file report
        :param file: the file name (relative or absolute) of the execution report.
        :return: None
        """
        self.document.add_heading("Last Execution report", 1)
        reporter = {}
        failed_count = 0
        succeed_count = 0
        test_count = 0
        current_feature = None
        current_scenario = None
        last_status = "skipped"
        with open(file) as report:
            for line in report.readlines():
                if re.match("Feature.*", line):
                    if current_feature is not None:
                        reporter[current_feature].update({current_scenario: last_status})
                        if last_status == "failed":
                            failed_count += 1
                        elif last_status == "passed":
                            succeed_count += 1
                        test_count += 1
                        last_status = "skipped"
                        self.document.add_page_break()
                    current_feature = line.split(":")[1].lstrip(' ').rstrip()
                    reporter[current_feature] = {}
                    current_scenario = None  # No scenario for the current feature
                    print(reporter)
                    self.document.add_heading(line.rstrip(), 2)
                elif re.match(r"\s*Scenario.*", line):
                    if current_scenario is not None:
                        reporter[current_feature].update({current_scenario: last_status})
                        if last_status == "failed":
                            failed_count += 1
                        elif last_status == "passed":
                            succeed_count += 1
                        test_count += 1
                        last_status = "skipped"
                    current_scenario = line.split(":")[1].lstrip(' ').rstrip()
                    print(current_scenario)
                    self.document.add_heading(line.rstrip(), 3)
                elif re.match(".*passed.*", line):
                    last_status = "passed"
                    self.document.add_paragraph(line.rstrip(), style='No Spacing')
                elif re.match(".*failed.*", line):
                    last_status = "failed"
                    self.document.add_paragraph(line.rstrip(), style='No Spacing')
                else:
                    self.document.add_paragraph(line.rstrip(), style='No Spacing')

        self.document.add_page_break()
        self.document.add_heading("Last Execution summary", 1)
        table_instance = self.document.add_table(rows=1, cols=3, style='Light List Accent 3')
        header_cells = table_instance.rows[0].cells
        header_cells[0].text = "Feature"
        header_cells[1].text = "Scenario"
        header_cells[2].text = "Status"
        print(reporter)
        feature_keys = list(reporter.keys())
        feature_keys.sort()
        for feature_key in feature_keys:
            scenario_keys = list(reporter[feature_key].keys())
            scenario_keys.sort()
            print(scenario_keys)
            for scenario_key in scenario_keys:
                row_cells = table_instance.add_row().cells
                row_cells[0].text = feature_key
                row_cells[1].text = scenario_key
                row_cells[2].text = reporter[feature_key][scenario_key]


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument("--tag")
    parser.add_argument("--title")
    parser.add_argument("--repository")
    parser.add_argument("--output")
    parser.add_argument("--execution")

    args = parser.parse_args()

    if all([item is None for item in vars(args).values()]):
        print("all none")
    else:
        print("some not none")
