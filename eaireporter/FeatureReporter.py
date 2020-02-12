# -*- coding: utf-8 -*-
import re
from behave.parser import parse_file
from docx import Document
import glob
import logging

log = logging.getLogger(__name__)


class ExportUtilities:

    def __init__(self, feature_repository: str = None,
                 user_story_tag_prefix: str = None,
                 report_title: str = None):
        self.__feature_repository = feature_repository
        self.__user_story_tag_prefix = user_story_tag_prefix
        self.__report_title = report_title

    def create_application_documentation(self, report_file=None, output_file_name="demo.docx"):
        """
        Create a document (docx) object and read first all ".feature" files and
        add their contents into the document.

        If a report file name is provided, it will add a "last execution" section containing
        the data found in the report.

        The report file must be the "plain" report output file generated from behave.

        :param report_file: The report file path (absolute or relative)
        :param output_file_name : The exported file name by default "demo.docx"
        :return: None
        """

        document = Document()
        document.add_heading("{}".format(self.__report_title), 0)  # Document title
        document.add_page_break()
        for file in glob.iglob("{}/**/*.feature".format(self.__feature_repository),
                               recursive=True):  # Use the iterator as it's cleaner
            # log.debug()
            test = parse_file(file)  # Use the Behave parser in order to read the feature file
            self.add_heading(document=document, feature=test)
            ExportUtilities.add_description(document=document, feature=test)
            ExportUtilities.add_background(document=document, feature=test)
            ExportUtilities.add_scenario(document=document, feature=test)
            document.add_page_break()
        if report_file is not None:
            ExportUtilities.add_report(document=document, file=report_file)
        document.save(output_file_name)

    def add_heading(self, document=None, feature=None):
        """
        Add a the feature name as top level section
        :param document: the document object
        :param feature: the feature object
        :return:
        """
        document.add_heading(feature.name, 1)
        paragraph = document.add_paragraph("")
        matcher = [elem for elem in feature.tags if self.__user_story_tag_prefix in elem]
        paragraph.add_run("Related to the user story: {}".format(str(matcher).strip('[]')))

    @staticmethod
    def add_description(document=None, feature=None):
        """
        Add the feature description into the document.
        If the feature file line contains a "*" character it will create a bullet list.
        If the line contains the sentence "business rules" then it will print it in bold.
        Otherwise it will print the line as-is.
        :param document: the document object
        :param feature: the feature object
        :return: None
        """
        for line in feature.description:
            if re.match(r'\*.*', line):
                document.add_paragraph(line[1:], style='List Bullet')
            elif re.match('[Bb]usiness [Rr]ules.*', line):
                paragraph = document.add_paragraph("")
                paragraph.add_run(line).bold = True
            else:
                document.add_paragraph(line)

    @staticmethod
    def add_background(document=None, feature=None):
        """
        Add the background of each scenario into the document
        :param document: the document object
        :param feature: the feature object from where to retrieve the background
        :return: None
        """
        if feature.background is not None:
            ExportUtilities.print_scenario_title(document=document,
                                                 scenario_keyword=feature.background.keyword,
                                                 scenario_name=feature.background.name)
            ExportUtilities.print_steps(document=document, steps=feature.background.steps)

    @staticmethod
    def add_scenario(document=None, feature=None):
        """
        Add in the document all the scenarios attached to a feature.
        :param document: the document object to add into
        :param feature: the feature object
        :return: None
        """
        if feature.scenarios is not None:
            for scenario in feature.scenarios:
                ExportUtilities.print_scenario_title(document=document,
                                                     scenario_keyword=scenario.keyword,
                                                     scenario_name=scenario.name)
                ExportUtilities.print_steps(document=document, steps=scenario.steps)
                if scenario.type == 'scenario_outline':
                    ExportUtilities.print_examples(document=document, examples=scenario.examples)

    @staticmethod
    def print_examples(document=None, examples=None):
        """
        Add an example section for each example attached to a scenario outline
        :param document: the document object to add into
        :param examples: an example list
        :return: None
        """
        for example in examples:
            ExportUtilities.print_scenario_title(document=document,
                                                 scenario_keyword=example.keyword,
                                                 scenario_name=example.name, level=3)
            ExportUtilities.print_table(document=document, table=example.table)

    @staticmethod
    def print_scenario_title(document=None, scenario_keyword=None, scenario_name=None, level=2):
        """
        Print a section title in the document in the format keyword : name with a level 2
        :param document: the document object to add the section title
        :param scenario_keyword: the keyword such as Scenario or Scenario Outline or Example
        :param scenario_name: the scenario name
        :param level: the section level. By default level 2
        :return: None
        """
        document.add_heading("{}: {}".format(scenario_keyword, scenario_name), level=level)
        # paragraph = document.add_paragraph("")
        # paragraph.add_run("{}:".format(scenario_keyword)).bold = True
        # paragraph.add_run(scenario_name)

    @staticmethod
    def print_steps(document=None, steps=None):
        """
        Add a step section into the document.
        :param document: the document object to write in
        :param steps: the feature steps table
        :return: None
        """
        step_done = []
        for step in steps:
            if step.keyword in step_done:
                keyword = "And"
            else:
                step_done.append(step.keyword)
                keyword = step.keyword

            paragraph = document.add_paragraph("", style='No Spacing')
            paragraph.add_run(keyword).bold = True
            paragraph.add_run(" {}".format(step.name))

            if step.table is not None:
                ExportUtilities.print_table(document=document, table=step.table)

    @staticmethod
    def print_table(document=None, table=None):
        """
        Add a tabular word object in the document based on the feature table object
        :param document: the document object to insert the section
        :param table: a feature table object
        :return: None
        """
        number_of_column = len(table.headings)

        table_instance = document.add_table(rows=1,
                                            cols=number_of_column,
                                            style='Light List Accent 3')
        header_cells = table_instance.rows[0].cells
        for count, text in enumerate(table.headings):
            header_cells[count].text = str(text)

        for row in table.rows:
            row_cells = table_instance.add_row().cells
            for count, cell in enumerate(row.cells):
                row_cells[count].text = str(cell)

    @staticmethod
    def add_report(document=None, file=None):
        """
        Add a last execution section to a document. It reads an execution plain file report
        :param document: the document object to insert the section
        :param file: the file name (relative or absolute) of the execution report.
        :return: None
        """
        document.add_heading("Last Execution report", 1)
        reporter = {}
        failed_count = 0
        succeed_count = 0
        test_count = 0
        current_feature = None
        current_scenario = None
        last_status = "skipped"
        with open(file) as report:
            for line in report.readlines():
                if re.match("Feature.*", line):
                    if current_feature is not None:
                        reporter[current_feature].update({current_scenario: last_status})
                        if last_status == "failed":
                            failed_count += 1
                        elif last_status == "passed":
                            succeed_count += 1
                        test_count += 1
                        last_status = "skipped"
                        document.add_page_break()
                    current_feature = line.split(":")[1].lstrip(' ').rstrip()
                    reporter[current_feature] = {}
                    current_scenario = None  # No scenario for the current feature
                    print(reporter)
                    document.add_heading(line.rstrip(), 2)
                elif re.match(r"\s*Scenario.*", line):
                    if current_scenario is not None:
                        reporter[current_feature].update({current_scenario: last_status})
                        if last_status == "failed":
                            failed_count += 1
                        elif last_status == "passed":
                            succeed_count += 1
                        test_count += 1
                        last_status = "skipped"
                    current_scenario = line.split(":")[1].lstrip(' ').rstrip()
                    print(current_scenario)
                    document.add_heading(line.rstrip(), 3)
                elif re.match(".*passed.*", line):
                    last_status = "passed"
                    document.add_paragraph(line.rstrip(), style='No Spacing')
                elif re.match(".*failed.*", line):
                    last_status = "failed"
                    document.add_paragraph(line.rstrip(), style='No Spacing')
                else:
                    document.add_paragraph(line.rstrip(), style='No Spacing')

        document.add_page_break()
        document.add_heading("Last Execution summary", 1)
        table_instance = document.add_table(rows=1, cols=3, style='Light List Accent 3')
        header_cells = table_instance.rows[0].cells
        header_cells[0].text = "Feature"
        header_cells[1].text = "Scenario"
        header_cells[2].text = "Status"
        print(reporter)
        feature_keys = list(reporter.keys())
        feature_keys.sort()
        for feature_key in feature_keys:
            scenario_keys = list(reporter[feature_key].keys())
            scenario_keys.sort()
            print(scenario_keys)
            for scenario_key in scenario_keys:
                row_cells = table_instance.add_row().cells
                row_cells[0].text = feature_key
                row_cells[1].text = scenario_key
                row_cells[2].text = reporter[feature_key][scenario_key]
