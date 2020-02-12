# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Cm
import glob
import os
import logging
from docx.shared import Pt
from reporter.FeatureReporter import ExportUtilities

log = logging.getLogger(__name__)


class ScenarioEvidence:

    @staticmethod
    def create_evidence(png_folder=None,
                        destination_folder=None,
                        file_name=None,
                        list_history=None,
                        scenario=None):
        log.info("Creating evidence")
        log.debug("Received\n png_folder '{}',\n destination_folder '{}'"
                  "\n file_name '{}'"
                  "\n list_history '{}'"
                  "\n scenario '{}'".format(png_folder,
                                            destination_folder,
                                            file_name,
                                            list_history,
                                            scenario))
        # A document will be created
        document = Document()
        # Add the scenario in the top

        if scenario is not None and scenario:
            document.add_heading("Feature: {}".format(scenario.feature.name), level=1)
            tags = ', '.join(scenario.feature.tags)
            document.add_paragraph("Tags: {}\n".format(tags))
            ExportUtilities.add_description(document=document, feature=scenario.feature)
            document.add_paragraph("")
            # ExportUtilities.add_background(document=document, feature=scenario.feature)
            document.add_heading("Scenario: {}".format(scenario.name), level=2)
            ExportUtilities.print_steps(document=document, steps=scenario.steps)
            document.add_paragraph("")

        # Check the list_history for API request
        if list_history is not None and list_history:
            requests = []
            log.debug("Request attribute initialization")
            # Retrieve data from the list
            for item in list_history:

                method = item.request.method
                path_url = item.request.path_url
                url = item.request.url
                status_code = item.status_code
                body = item.request.body
                text = item.text

                if body is not None and not isinstance(body, str):
                    body = body.decode()

                log.debug("body:{}".format(str(body)))
                headers = item.request.headers

                log.debug("'\n' request_url: {} '\n' request_method: {}"
                          "'\n' request_path: {}'\n' request_status_code: {}"
                          "'\n' request_header: {} '\n' request_body: {}"
                          "'\n' request_text: {}".format(method,
                                                         path_url,
                                                         url,
                                                         status_code,
                                                         headers,
                                                         body,
                                                         text))

                request = ["Endpoint: {}".format(url),
                           "Method: {}".format(method),
                           "Path_url: {}".format(path_url),
                           "Status_code: {}".format(status_code),
                           "Headers: {}".format(headers),
                           "Body: {}".format(body),
                           "Response: {}".format(text)]
                # All requests are added to the list
                requests.append(request)
            # Create the evidence
            for i, items in enumerate(requests):
                for request in items:
                    # find title & request
                    pos = request.find(":") + 1
                    title = request[0:pos]
                    request = request[pos:len(request)]

                    if "Endpoint:" in request:
                        # Header style & title
                        paragraph_header = document.add_paragraph(
                            "{0}".format(title))
                        paragraph_header.style = document.styles['Heading 3']
                        font = paragraph_header.style.font
                        font.size = Pt(12)
                        p_format = paragraph_header.style.paragraph_format
                        p_format.left_indent = Pt(-70)
                        p_format.right_indent = Pt(-90)
                        text = paragraph_header.add_run(
                            "{0}{1}{1}".format(request, '\n'))
                        text.style = 'Emphasis'
                        continue
                    # Style for the rest of the document
                    paragraph_elements = document.add_paragraph(
                        "{}".format(title))
                    font = paragraph_elements.style.font
                    font.bold = True
                    text = paragraph_elements.add_run(request)
                    text.bold = False
                # Add a page break between requests except for the last one
                if i == len(requests) - 1:
                    continue
                document.add_page_break()

        # Check if png must be added in the evidence
        if png_folder is not None and png_folder:
            search_sequence = "{}/*.png".format(png_folder)
            log.debug("Search sequence '{}'".format(search_sequence))
            files = glob.glob(search_sequence)
            files.sort(key=os.path.getmtime)

            log.debug("Retrieved '{}' file(s)".format(len(files)))

            for file in files:
                document.add_picture(file, width=Cm(18))
                document.add_paragraph("")

            failure = glob.glob("{}/*.txt".format(png_folder))
            if failure:
                with open(failure[0]) as failure_file:
                    document.add_heading("Failure cause message", level=1)
                    document.add_paragraph(failure_file.read())

        document.save(os.path.join(destination_folder, file_name))
